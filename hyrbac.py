#!/usr/bin/env python3
# @paciente23256

import json
from docx import Document

# Ler JSON
with open("politica.json", "r", encoding="utf-8") as f:
data = json.load(f)


doc = Document()


# Capa
doc.add_heading(data["capa"]["titulo"], level=0)
doc.add_paragraph(data["capa"]["subtitulo"])
doc.add_paragraph(f"Data: {data['capa']['data']}")
doc.add_paragraph(f"Versão: {data['capa']['versao']}")
doc.add_paragraph(f"Responsável: {data['capa']['responsavel']}")


# Objetivos
doc.add_heading("Objetivos", level=1)
for obj in data["objetivos"]:
doc.add_paragraph(f"- {obj}")


# Princípios
doc.add_heading("Princípios da Política", level=1)
for p in data["principios"]:
doc.add_paragraph(f"- {p}")


# Matriz RBAC
doc.add_heading("Matriz de Funções RBAC", level=1)
table = doc.add_table(rows=1, cols=5)
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text = ["Função", "Âmbito", "Permissões", "Perfil", "Notas"]


for role in data["roles"]:
row = table.add_row().cells
row[0].text = role["funcao"]
row[1].text = role["ambito"]
row[2].text = role["permissoes"]
row[3].text = role["perfil"]
row[4].text = role["notas"]


# Matriz SoD
doc.add_heading("Matriz de Segregação de Funções (SoD)", level=1)
sod_table = doc.add_table(rows=1, cols=2)
sod_hdr = sod_table.rows[0].cells
sod_hdr[0].text, sod_hdr[1].text = ["Funções em Conflito", "Motivo"]


for sod in data["sod"]:
row = sod_table.add_row().cells
row[0].text = sod["funcoes"]
row[1].text = sod["motivo"]


# Salvaguardas
doc.add_heading("Salvaguardas Adicionais", level=1)
for s in data["salvaguardas"]:
doc.add_paragraph(f"- {s}")


# Guardar documento
doc.save("Politica_IAM_AD_Hibrido.docx")
print("Documento Word criado com sucesso: Politica_IAM_AD_Hibrido.docx")
